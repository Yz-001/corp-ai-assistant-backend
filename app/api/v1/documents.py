"""Documents API endpoints."""

import os
import asyncio
from pathlib import Path

from fastapi import APIRouter, HTTPException, Query, UploadFile, File, Form, BackgroundTasks
from sqlalchemy import select, func

from app.api.deps import DBSession, CurrentUser
from app.core.config import settings
from app.models.document import Document, DocumentChunk
from app.schemas import (
    BaseResponse,
    PaginatedResponse,
    DocumentUploadResponse,
    DocumentResponse,
    DocumentListResponse,
    DocumentChunkResponse,
)
from app.utils.id import generate_id

router = APIRouter()

# Upload directory
UPLOAD_DIR = Path("upload")
UPLOAD_DIR.mkdir(exist_ok=True)


def get_file_extension(content_type: str) -> str:
    """Get file extension from content type."""
    extensions = {
        "application/pdf": ".pdf",
        "application/msword": ".doc",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": ".docx",
        "text/plain": ".txt",
        "text/markdown": ".md",
    }
    return extensions.get(content_type, "")


async def process_document(doc_id: str, file_path: Path, tenant_id: str):
    """Process document: extract text, split into chunks, and store."""
    from app.core.database import async_session_maker
    
    print(f"[Document Processing] Starting processing for doc: {doc_id}")
    
    async with async_session_maker() as db:
        try:
            # Get document
            result = await db.execute(select(Document).where(Document.id == doc_id))
            doc = result.scalar_one_or_none()
            if not doc:
                print(f"[Document Processing] Document not found: {doc_id}")
                return
            
            print(f"[Document Processing] Found document: {doc.name}")
            
            # Update status to processing
            doc.status = "processing"
            await db.commit()
            print(f"[Document Processing] Status updated to processing")
            
            # Extract text based on file type
            text_content = ""
            file_ext = file_path.suffix.lower()
            print(f"[Document Processing] File extension: {file_ext}")
            
            if file_ext == ".txt" or file_ext == ".md":
                with open(file_path, "r", encoding="utf-8") as f:
                    text_content = f.read()
                print(f"[Document Processing] Read text file, length: {len(text_content)}")
            elif file_ext == ".pdf":
                try:
                    import pypdf
                    with open(file_path, "rb") as f:
                        reader = pypdf.PdfReader(f)
                        for page in reader.pages:
                            text_content += page.extract_text() + "\n"
                    print(f"[Document Processing] Extracted PDF, length: {len(text_content)}")
                except ImportError as e:
                    print(f"[Document Processing] pypdf not installed: {e}")
                    text_content = ""
            elif file_ext in [".doc", ".docx"]:
                try:
                    from docx import Document as DocxDocument
                    docx_doc = DocxDocument(file_path)
                    for para in docx_doc.paragraphs:
                        text_content += para.text + "\n"
                    print(f"[Document Processing] Extracted DOCX, length: {len(text_content)}")
                except ImportError as e:
                    print(f"[Document Processing] python-docx not installed: {e}")
                    text_content = ""
            
            if not text_content.strip():
                print(f"[Document Processing] No text content extracted")
                doc.status = "failed"
                doc.error_message = "无法提取文档内容"
                await db.commit()
                return
            
            # Split into chunks (simple implementation - by paragraphs/sections)
            chunk_size = 500  # characters per chunk
            overlap = 50  # overlap between chunks
            
            chunks = []
            text = text_content.strip()
            
            # Simple chunking by character count with overlap
            start = 0
            chunk_index = 0
            while start < len(text):
                end = start + chunk_size
                chunk_text = text[start:end]
                
                # Try to end at a sentence or paragraph boundary
                if end < len(text):
                    # Look for sentence boundary
                    last_period = chunk_text.rfind("。")
                    last_newline = chunk_text.rfind("\n")
                    boundary = max(last_period, last_newline)
                    if boundary > chunk_size // 2:
                        chunk_text = chunk_text[:boundary + 1]
                        end = start + boundary + 1
                
                if chunk_text.strip():
                    # Calculate approximate token count (rough estimate: 1 token ≈ 2 Chinese chars)
                    token_count = len(chunk_text) // 2
                    
                    chunk = DocumentChunk(
                        id=generate_id(),
                        tenant_id=tenant_id,
                        document_id=doc_id,
                        chunk_index=chunk_index,
                        content=chunk_text.strip(),
                        token_count=token_count,
                        chunk_metadata={"char_count": len(chunk_text)},
                    )
                    chunks.append(chunk)
                    chunk_index += 1
                
                start = end - overlap if end < len(text) else end
            
            print(f"[Document Processing] Created {len(chunks)} chunks")
            
            # Save chunks
            for chunk in chunks:
                db.add(chunk)
            
            # Update document status
            doc.status = "completed"
            doc.chunk_count = len(chunks)
            doc.file_size = os.path.getsize(file_path)
            await db.commit()
            print(f"[Document Processing] Document processing completed successfully")
            
        except Exception as e:
            print(f"[Document Processing] Error: {e}")
            import traceback
            traceback.print_exc()
            # Update document status to failed
            try:
                result = await db.execute(select(Document).where(Document.id == doc_id))
                doc = result.scalar_one_or_none()
                if doc:
                    doc.status = "failed"
                    doc.error_message = str(e)
                    await db.commit()
            except Exception as e2:
                print(f"[Document Processing] Failed to update error status: {e2}")


@router.post("/upload", response_model=BaseResponse[DocumentUploadResponse])
async def upload_document(
    background_tasks: BackgroundTasks,
    current_user: CurrentUser,
    db: DBSession,
    file: UploadFile = File(...),
    visibility: str = Form(default="private"),
    remark: str | None = Form(default=None),
):
    """Upload a document."""
    # Validate file type
    allowed_types = ["application/pdf", "application/msword", 
                     "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                     "text/plain", "text/markdown"]
    if file.content_type not in allowed_types:
        raise HTTPException(status_code=400, detail="不支持的文件类型，仅支持 PDF、Word、TXT、Markdown 格式")
    
    # Read file content
    content = await file.read()
    file_size = len(content)
    
    # Generate storage path
    file_ext = get_file_extension(file.content_type or "")
    storage_name = f"{generate_id()}{file_ext}"
    storage_path = UPLOAD_DIR / storage_name
    
    # Save file
    with open(storage_path, "wb") as f:
        f.write(content)
    
    # Create document record
    doc = Document(
        id=generate_id(),
        tenant_id=current_user.tenant_id,
        name=file.filename or "unnamed",
        file_name=file.filename or "unnamed",
        file_type=file.content_type or "unknown",
        file_size=file_size,
        storage_path=str(storage_path),
        visibility=visibility,
        status="pending",
        created_by=current_user.id,
    )
    db.add(doc)
    await db.commit()
    await db.refresh(doc)
    
    # Start background processing
    background_tasks.add_task(
        process_document, 
        doc.id, 
        storage_path, 
        current_user.tenant_id
    )
    
    return BaseResponse(
        data=DocumentUploadResponse(
            documentId=doc.id,
            status=doc.status,
        ),
        message="文档上传成功，正在处理中..."
    )


@router.get("", response_model=BaseResponse[PaginatedResponse[DocumentListResponse]])
async def get_documents(
    current_user: CurrentUser,
    db: DBSession,
    pageNum: int = Query(1, ge=1),
    pageSize: int = Query(20, ge=1, le=100),
    keyword: str | None = Query(None),
    status: str | None = Query(None),
    fileType: str | None = Query(None),
):
    """Get documents list."""
    query = select(Document).where(Document.tenant_id == current_user.tenant_id)
    
    if keyword:
        query = query.where(Document.name.ilike(f"%{keyword}%"))
    if status:
        query = query.where(Document.status == status)
    if fileType:
        query = query.where(Document.file_type == fileType)
    
    count_query = select(func.count()).select_from(query.subquery())
    total_result = await db.execute(count_query)
    total = total_result.scalar()
    
    query = query.order_by(Document.created_at.desc())
    query = query.offset((pageNum - 1) * pageSize).limit(pageSize)
    
    result = await db.execute(query)
    docs = result.scalars().all()
    
    items = [
        DocumentListResponse(
            documentId=d.id,
            name=d.name,
            fileName=d.file_name,
            fileType=d.file_type,
            fileSize=d.file_size,
            tenantName=None,
            visibility=d.visibility,
            status=d.status,
            chunkCount=d.chunk_count,
            createdAt=d.created_at,
        )
        for d in docs
    ]
    
    return BaseResponse(
        data=PaginatedResponse(
            items=items,
            total=total,
            pageNum=pageNum,
            pageSize=pageSize,
        )
    )


@router.get("/{documentId}", response_model=BaseResponse[DocumentResponse])
async def get_document(documentId: str, current_user: CurrentUser, db: DBSession):
    """Get document details."""
    result = await db.execute(
        select(Document).where(
            Document.id == documentId,
            Document.tenant_id == current_user.tenant_id,
        )
    )
    doc = result.scalar_one_or_none()
    
    if doc is None:
        raise HTTPException(status_code=404, detail="Document not found")
    
    return BaseResponse(
        data=DocumentResponse(
            documentId=doc.id,
            name=doc.name,
            fileName=doc.file_name,
            fileType=doc.file_type,
            fileSize=doc.file_size,
            tenantId=doc.tenant_id,
            tenantName=None,
            visibility=doc.visibility,
            status=doc.status,
            chunkCount=doc.chunk_count,
            createdBy=doc.created_by,
            errorMessage=doc.error_message,
            createdAt=doc.created_at,
            updatedAt=doc.updated_at,
        )
    )


@router.delete("/{documentId}", response_model=BaseResponse)
async def delete_document(documentId: str, current_user: CurrentUser, db: DBSession):
    """Delete a document."""
    result = await db.execute(
        select(Document).where(
            Document.id == documentId,
            Document.tenant_id == current_user.tenant_id,
        )
    )
    doc = result.scalar_one_or_none()
    
    if doc is None:
        raise HTTPException(status_code=404, detail="Document not found")
    
    await db.delete(doc)
    await db.commit()
    
    return BaseResponse(message="Document deleted successfully")


@router.post("/{documentId}/retry", response_model=BaseResponse)
async def retry_document(documentId: str, current_user: CurrentUser, db: DBSession):
    """Retry document processing."""
    result = await db.execute(
        select(Document).where(
            Document.id == documentId,
            Document.tenant_id == current_user.tenant_id,
        )
    )
    doc = result.scalar_one_or_none()
    
    if doc is None:
        raise HTTPException(status_code=404, detail="Document not found")
    
    if doc.status != "failed":
        raise HTTPException(status_code=400, detail="Document is not in failed status")
    
    doc.status = "pending"
    doc.error_message = None
    await db.commit()
    
    # TODO: Trigger async processing
    
    return BaseResponse(message="Document retry initiated")


@router.get("/{documentId}/chunks", response_model=BaseResponse[PaginatedResponse[DocumentChunkResponse]])
async def get_document_chunks(
    documentId: str,
    current_user: CurrentUser,
    db: DBSession,
    pageNum: int = Query(1, ge=1),
    pageSize: int = Query(20, ge=1, le=100),
):
    """Get document chunks."""
    result = await db.execute(
        select(Document).where(
            Document.id == documentId,
            Document.tenant_id == current_user.tenant_id,
        )
    )
    doc = result.scalar_one_or_none()
    
    if doc is None:
        raise HTTPException(status_code=404, detail="Document not found")
    
    query = select(DocumentChunk).where(DocumentChunk.document_id == documentId)
    
    count_query = select(func.count()).select_from(query.subquery())
    total_result = await db.execute(count_query)
    total = total_result.scalar()
    
    query = query.order_by(DocumentChunk.chunk_index.asc())
    query = query.offset((pageNum - 1) * pageSize).limit(pageSize)
    
    result = await db.execute(query)
    chunks = result.scalars().all()
    
    items = [
        DocumentChunkResponse(
            chunkId=c.id,
            chunkIndex=c.chunk_index,
            content=c.content,
            tokenCount=c.token_count,
            metadata=c.chunk_metadata,
        )
        for c in chunks
    ]
    
    return BaseResponse(
        data=PaginatedResponse(
            items=items,
            total=total,
            pageNum=pageNum,
            pageSize=pageSize,
        )
    )
